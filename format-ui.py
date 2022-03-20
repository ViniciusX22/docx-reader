from docx import Document
from docx.shared import Pt, Cm, Length
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
import csv
from tkinter import Tk
from tkinter.filedialog import askopenfilename, askdirectory

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)


ap = argparse.ArgumentParser()
ap.add_argument("-ls", '--line-spacing', default=1.5,
                type=float, help="Line spacing of each paragraph")
ap.add_argument('-ff', '--font-family', default='Arial',
                help="Font family of document")
ap.add_argument('-fs', '--font-size', default=Pt(12),
                help="Font size of document")
ap.add_argument('-m', '--margins', nargs=4,
                default=[3, 2, 2, 3], help="Margins of each page")
ap.add_argument('-a', '--alignment', default='justify',
                choices=['justify', 'left', 'center', 'right'], help="Margins of each page")


def add_cover_template(doc, values):
    template = Document(resource_path('cover-template.docx'))

    for i in range(len(template.paragraphs)):
        for run in template.paragraphs[i].runs:
            for field in values:
                if 'professor' in field and field in run.text and not values[field]:
                    template.paragraphs[i].clear()
                    template.paragraphs[i-1].clear()
                    template.paragraphs[i+1].clear()
                    break
                run.text = run.text.replace(f'{{{field}}}', values[field])
                if not run.text:
                    run.clear()

    template.add_section()
    for p in doc.paragraphs:
        new_p = template.add_paragraph('')
        for run in p.runs:
            new_run = new_p.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.all_caps = run.font.all_caps
            new_run.font.bold = run.font.bold
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.color.theme_color = run.font.color.theme_color
            new_run.font.complex_script = run.font.complex_script
            new_run.font.cs_bold = run.font.cs_bold
            new_run.font.cs_italic = run.font.cs_italic
            new_run.font.double_strike = run.font.double_strike
            new_run.font.emboss = run.font.emboss
            new_run.font.hidden = run.font.hidden
            new_run.font.highlight_color = run.font.highlight_color
            new_run.font.imprint = run.font.imprint
            new_run.font.italic = run.font.italic
            new_run.font.math = run.font.math
            new_run.font.name = run.font.name
            new_run.font.no_proof = run.font.no_proof
            new_run.font.outline = run.font.outline
            new_run.font.rtl = run.font.rtl
            new_run.font.shadow = run.font.shadow
            new_run.font.size = run.font.size
            new_run.font.small_caps = run.font.small_caps
            new_run.font.snap_to_grid = run.font.snap_to_grid
            new_run.font.spec_vanish = run.font.spec_vanish
            new_run.font.strike = run.font.strike
            new_run.font.subscript = run.font.subscript
            new_run.font.superscript = run.font.superscript
            new_run.font.underline = run.font.underline
            new_run.font.web_hidden = run.font.web_hidden
            new_run.style = run.style
        new_p.alignment = p.alignment
        new_p.style = p.style
        new_p.paragraph_format.alignment = p.paragraph_format.alignment
        new_p.paragraph_format.first_line_indent = p.paragraph_format.first_line_indent
        new_p.paragraph_format.keep_together = p.paragraph_format.keep_together
        new_p.paragraph_format.keep_with_next = p.paragraph_format.keep_with_next
        new_p.paragraph_format.left_indent = p.paragraph_format.left_indent
        new_p.paragraph_format.line_spacing = p.paragraph_format.line_spacing
        new_p.paragraph_format.page_break_before = p.paragraph_format.page_break_before
        new_p.paragraph_format.right_indent = p.paragraph_format.right_indent
        new_p.paragraph_format.space_after = p.paragraph_format.space_after
        new_p.paragraph_format.space_before = p.paragraph_format.space_before
        new_p.paragraph_format.widow_control = p.paragraph_format.widow_control
        for ts in p.paragraph_format.tab_stops:
            new_p.paragraph_format.add_tab_stop(
                ts.position, ts.alignment, ts.leader)

    return template


args = vars(ap.parse_args())
cover_values = None

Tk().withdraw()
# show an "Open" dialog box and return the path to the selected file
filename = askopenfilename(
    title="Abrir documento a ser formatado", initialfile="documento.docx")
cover_file = askopenfilename(
    title="Abrir valores para capa", initialfile="capa.csv")

if not filename:
    print('Document is required')
    exit(1)

if cover_file:
    with open(cover_file, newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        values = {}
        for row in reader:
            values["instituition"] = row['Instiuicao']
            values["students"] = row['Estudantes']
            values["title"] = row['Titulo']
            values["subtitle"] = row['Subtitulo']
            values["city"] = row['Cidade']
            values["state"] = row['Estado']
            values["day"] = row['Dia']
            values["month"] = row['Mes']
            values["year"] = row['Ano']
            values["group"] = row['Grupo']
            values["description"] = row['Descricao']
            values["board"] = row['Banca Eximinadora']
            values["professor1"] = row['Professor 1']
            values["professor2"] = row['Professor 2']
            values["professor3"] = row['Professor 3']

        cover_values = values

doc = Document(filename)

for s in doc.sections:
    s.top_margin = Cm(args['margins'][0])
    s.right_margin = Cm(args['margins'][1])
    s.bottom_margin = Cm(args['margins'][2])
    s.left_margin = Cm(args['margins'][3])

for p in doc.paragraphs:
    p.paragraph_format.line_spacing = args['line_spacing']
    p.paragraph_format.alignment = eval(f'WD_ALIGN_PARAGRAPH.{args["alignment"].upper()}') if p.style and p.style.name not in [
        'Title', 'Subtitle'] else WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = args['font_family']
        r.font.size = args['font_size'] if p.style and p.style.name == 'Normal' else args['font_size'] + \
            Pt(2)

if cover_values:
    doc = add_cover_template(doc, cover_values)

doc.save(filename[:filename.index('.')] +
         '-formated' + filename[filename.index('.'):])
